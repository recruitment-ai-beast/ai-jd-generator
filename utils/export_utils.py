"""
Export utilities — PDF, DOCX, TXT.
All export logic isolated here.
"""

import logging
from io import BytesIO
from typing import Optional

logger = logging.getLogger(__name__)


def export_as_txt(content: str) -> bytes:
    """Export JD as plain text."""
    return content.encode("utf-8")


def export_as_pdf(content: str, job_title: str) -> Optional[bytes]:
    """Export JD as PDF using fpdf2."""
    try:
        from fpdf import FPDF

        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, job_title, ln=True, align="C")
        pdf.ln(5)
        pdf.set_font("Helvetica", size=11)

        for line in content.split("\n"):
            line = line.strip()
            if not line:
                pdf.ln(4)
            elif line.startswith("##"):
                pdf.set_font("Helvetica", "B", 13)
                pdf.multi_cell(0, 8, line.replace("#", "").strip())
                pdf.set_font("Helvetica", size=11)
            else:
                pdf.multi_cell(0, 7, line)

        return pdf.output()

    except Exception as e:
        logger.error(f"PDF export failed: {e}")
        return None


def export_as_docx(content: str, job_title: str) -> Optional[bytes]:
    """Export JD as DOCX using python-docx."""
    try:
        from docx import Document

        doc = Document()
        doc.add_heading(job_title, 0)

        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue
            elif line.startswith("##"):
                doc.add_heading(line.replace("#", "").strip(), level=2)
            elif line.startswith("-") or line.startswith("•"):
                doc.add_paragraph(line, style="List Bullet")
            else:
                doc.add_paragraph(line)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    except Exception as e:
        logger.error(f"DOCX export failed: {e}")
        return None
